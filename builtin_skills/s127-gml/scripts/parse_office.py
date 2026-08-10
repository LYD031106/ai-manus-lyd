#!/usr/bin/env python3
"""PDF / Word / Excel / CSV → Markdown（沙箱内本地解析，文本部分不消耗 token）。

    python3 $S/parse_office.py 通告.pdf 附件1.docx 坐标.csv -o 解析结果.md

多个文件的结果按文件名分节合并。

## 分工原则：文本走脚本，图走模型

**有文本层的文件，文本一律本地抽取** —— 文本层是无损的，过一遍 OCR 只会变差
（实测：模型会把原文 `宽32.3米、、满载吃水` 顺手改写成 `宽32.3m`，法规数据不能这样）。
**只有图里的信息才交给模型**，用 `--image-dir` 导出后交给 agent 的
`parse_regulation` 工具。

按文件类型分三路：

| 情况 | 本脚本做什么 | 还需要 parse_regulation 吗 |
|---|---|---|
| PDF 有文本层、无大图 | 抽文本 + 表格 | 不需要 |
| PDF 有文本层、有大图 | 抽文本 + 表格，把**含大图的页**渲染成 PNG 导出 | 需要，只读导出的 PNG |
| PDF 无文本层（扫描件） | 只报告页数，不抽文本 | 需要，整份文件交给它 |
| Word 有内嵌图 | 抽正文 + 表格，导出大图 | 需要，只读导出的图 |
| Word 无图 / Excel / CSV | 全部本地抽完 | 不需要 |

图片按尺寸过滤（默认 ≥150×150）：真实公文里混着大量页面装饰、印章碎片、
网页图标（珠江口那份就有 802 张小图），全发给模型既浪费又添噪。

PDF 内嵌图的编码五花八门（JBIG2 / CCITTFax / JPX / DCT），逐个解码要引一堆依赖，
所以 PDF 走**整页渲染**而不是抽取内嵌图 —— 渲染结果不受原图编码影响，
且保留了图在页面里的上下文（图题、图例、周边文字）。
"""
from __future__ import annotations

import argparse
import csv
import io
import sys
from pathlib import Path
from typing import Any, List, Tuple

# 低于这个尺寸的图视为装饰，不导出（单位：PDF 用点，Word 用像素）
MIN_IMAGE_SIDE = 150
# 判定扫描件的阈值：平均每页可抽取字符数低于此值 → 认为没有文本层
SCANNED_CHARS_PER_PAGE = 50
# 整页渲染倍率，2 ≈ 144 dpi，够模型读图又不至于让 token 爆掉
PAGE_RENDER_SCALE = 2

# 多模态 API 能直接读的位图格式；EMF/WMF 等矢量格式读不了，需要另行转换
API_READABLE = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
}
OTHER_IMAGE_EXT = {
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
    "image/emf": ".emf",
    "image/wmf": ".wmf",
    "image/tiff": ".tiff",
    "image/bmp": ".bmp",
}


def cell_text(value: Any) -> str:
    """单元格 → Markdown 安全文本：换行压成空格，竖线转义，否则表格会散。"""
    if value is None:
        return ""
    return str(value).replace("\r", " ").replace("\n", " ").replace("|", "\\|").strip()


def rows_to_markdown(rows: List[List[Any]]) -> str:
    """二维单元格 → Markdown 表格。按最宽的一行补齐列数。"""
    if not rows:
        return "(空文件)"
    width = max((len(row) for row in rows), default=0)
    if width == 0:
        return "(空文件)"
    padded = [[cell_text(c) for c in row] + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(padded[0]) + " |",
             "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in padded[1:])
    return "\n".join(lines)


def docx_table_to_markdown(table) -> str:
    """Word 表格 → Markdown，处理合并单元格。

    python-docx 会把跨列合并（gridSpan）的单元格在网格里重复展开，直接输出
    会得到「100 | 100 | 100」这种重复列。更麻烦的是同一张表各行的合并方式
    可能不同（真实公文里就有），展开后数值与表头对不齐。这里把合并的延续
    位置留空，并在某些行的合并结构与表头不一致时附一行提示 —— 那种表用
    平铺 Markdown 无法忠实表达，必须让下游知道不能照列硬读。

    纵向合并（vMerge）不用管：python-docx 已把延续行解析成合并起始格的文本。
    """
    rows: List[List[str]] = []
    patterns: List[tuple] = []
    for row in table.rows:
        cells: List[str] = []
        pattern: List[bool] = []
        prev = None
        for cell in row.cells:
            # 同一个 <w:tc> 出现在相邻网格位 → 跨列合并的延续位
            spanned = prev is not None and cell._tc is prev._tc
            cells.append("" if spanned else cell.text)
            pattern.append(spanned)
            prev = cell
        rows.append(cells)
        patterns.append(tuple(pattern))

    markdown = rows_to_markdown(rows)
    if any(True in p for p in patterns):
        odd = [i + 1 for i, p in enumerate(patterns) if p != patterns[0]]
        if odd:
            markdown += (
                "\n\n> ⚠️ 本表含跨列合并，且第 "
                + "、".join(map(str, odd))
                + " 行（含表头计）的合并方式与表头不同，这些行的列与表头可能对不齐；"
                "涉及坐标或数值时请对照原文核对。"
            )
    return markdown


def _image_side_ok(blob: bytes) -> bool:
    """按像素尺寸判断是否值得送模型；读不出尺寸时保守地当作值得。

    公文 Word 里常混着分享按钮、字号图标一类装饰图（实测有 2 KB 的），
    发给模型纯属浪费还添噪。
    """
    try:
        from PIL import Image
        with Image.open(io.BytesIO(blob)) as im:
            w, h = im.size
        return w >= MIN_IMAGE_SIDE and h >= MIN_IMAGE_SIDE
    except Exception:
        return True


def export_images(doc, stem: str, image_dir: Path) -> Tuple[List[Path], List[Path], int]:
    """导出 Word 内嵌图片，返回 (API 可读的, 其余格式的, 因过小而跳过的数量)。"""
    readable: List[Path] = []
    other: List[Path] = []
    skipped = 0
    n = 0
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
            ctype = (rel.target_part.content_type or "image/png").lower()
        except Exception:  # 个别损坏的图片关系不应影响整篇解析
            print("[警告] 跳过一张无法读取的内嵌图片", file=sys.stderr)
            continue
        if not _image_side_ok(blob):
            skipped += 1
            continue
        n += 1
        ext = API_READABLE.get(ctype) or OTHER_IMAGE_EXT.get(ctype) or ".bin"
        target = image_dir / f"{stem}_img{n}{ext}"
        target.write_bytes(blob)
        (readable if ctype in API_READABLE else other).append(target)
    return readable, other, skipped


def parse_docx(path: Path, image_dir: Path | None) -> str:
    """解析 .docx：段落（含标题层级）+ 表格；内嵌图片导出后交给 parse_regulation。"""
    from docx import Document

    doc = Document(str(path))
    lines: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading" in style:
            # 从样式名里取出标题级别，如 "Heading 2" → 2
            level = next((int(ch) for ch in style if ch.isdigit()), 1)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        if not table.rows:
            continue
        lines.append("")
        lines.append(docx_table_to_markdown(table))
        lines.append("")

    n_images = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    if n_images:
        lines.append("")
        if image_dir is None:
            lines.append(
                f"> ⚠️ 本文件含 {n_images} 张内嵌图片，未导出（本次未指定 --image-dir）。"
                f"若图中有坐标或范围信息，请加 --image-dir 重跑并用 parse_regulation 解析图片。"
            )
        else:
            readable, other, skipped = export_images(doc, path.stem, image_dir)
            if readable:
                lines.append(f"> 本文件含 {n_images} 张内嵌图片，已导出以下可解析图片，"
                             f"请用 parse_regulation 工具读取其中的坐标与范围标注：")
                lines.extend(f"> - {p}" for p in readable)
            if other:
                lines.append(f"> ⚠️ 另有 {len(other)} 张为矢量/非位图格式，多模态 API 读不了，"
                             f"需先转成 PNG，或改用该文件的 PDF 版本：")
                lines.extend(f"> - {p}" for p in other)
            if skipped:
                lines.append(f"> 另有 {skipped} 张小于 {MIN_IMAGE_SIDE}×{MIN_IMAGE_SIDE} 的"
                             f"装饰图（图标、按钮之类），已跳过不送模型。")
            if not readable and not other:
                lines.append(f"> 本文件的 {n_images} 张图片全为装饰图，无需调用 parse_regulation。")
            print(f"[图片] {path.name}：导出 {len(readable)} 张可解析"
                  + (f"、{len(other)} 张需转换" if other else "")
                  + (f"、跳过 {skipped} 张装饰图" if skipped else ""), file=sys.stderr)

    return "\n".join(lines)


def parse_csv(path: Path, _image_dir: Path | None = None) -> str:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    return rows_to_markdown(list(csv.reader(io.StringIO(text))))


def parse_excel(path: Path, _image_dir: Path | None = None) -> str:
    """逐个工作表解析 .xlsx，公式取计算结果。"""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        rows = [list(row) for row in ws.iter_rows(values_only=True)]
        if not rows:
            continue
        parts.append(f"### Sheet: {ws.title}\n{rows_to_markdown(rows)}")
    return "\n\n".join(parts) if parts else "(空表格)"


def _render_pages(path: Path, pages: List[int], stem: str, image_dir: Path) -> List[Path]:
    """把指定页（1 基）渲染成 PNG 导出，返回文件路径。

    渲染整页而非抽取内嵌图：PDF 内嵌图编码有 JBIG2 / CCITTFax / JPX 等多种，
    逐个解码要引一堆依赖；渲染不受编码影响，还保留了图题图例等上下文。
    """
    import pypdfium2 as pdfium

    out: List[Path] = []
    doc = pdfium.PdfDocument(str(path))
    try:
        for pno in pages:
            img = doc[pno - 1].render(scale=PAGE_RENDER_SCALE).to_pil()
            target = image_dir / f"{stem}_p{pno}.png"
            img.save(target, "PNG")
            out.append(target)
    finally:
        doc.close()
    return out


def parse_pdf(path: Path, image_dir: Path | None) -> str:
    """解析 PDF：有文本层就本地抽文本与表格，含大图的页另行渲染导出。

    扫描件（没有文本层）不在这里硬抽 —— 抽出来是空的或乱的，
    直接报告让调用方把整份文件交给 parse_regulation。
    """
    import pdfplumber

    lines: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        n_pages = len(pdf.pages)
        page_texts: List[str] = []
        page_tables: List[List[List[List[Any]]]] = []
        img_pages: List[int] = []
        for i, page in enumerate(pdf.pages, start=1):
            page_texts.append(page.extract_text() or "")
            page_tables.append(page.extract_tables() or [])
            if any(abs(im["x1"] - im["x0"]) >= MIN_IMAGE_SIDE
                   and abs(im["top"] - im["bottom"]) >= MIN_IMAGE_SIDE
                   for im in page.images):
                img_pages.append(i)

    total_chars = sum(len(t) for t in page_texts)
    scanned = total_chars / max(n_pages, 1) < SCANNED_CHARS_PER_PAGE

    if scanned:
        lines.append(f"> ⚠️ 本 PDF 共 {n_pages} 页，**没有文本层**（平均每页仅可抽取 "
                     f"{total_chars // max(n_pages, 1)} 字符），是扫描件。")
        lines.append("> 本脚本不做 OCR —— 请把**整份 PDF** 直接交给 parse_regulation 工具解析。")
        print(f"[扫描件] {path.name}：{n_pages} 页无文本层，需交 parse_regulation",
              file=sys.stderr)
        return "\n".join(lines)

    for i, (text, tables) in enumerate(zip(page_texts, page_tables), start=1):
        body = text.strip()
        if not body and not tables:
            continue
        lines.append(f"### 第 {i} 页")
        if body:
            lines.append(body)
        for t in tables:
            lines.append("")
            lines.append(rows_to_markdown(t))
        lines.append("")

    if img_pages:
        if image_dir is None:
            lines.append(f"> ⚠️ 第 {', '.join(map(str, img_pages))} 页含图件，未导出"
                         f"（本次未指定 --image-dir）。图中若有坐标或范围信息，"
                         f"请加 --image-dir 重跑并用 parse_regulation 解析。")
        else:
            rendered = _render_pages(path, img_pages, path.stem, image_dir)
            lines.append(f"> 第 {', '.join(map(str, img_pages))} 页含图件，"
                         f"已按页渲染导出 {len(rendered)} 张 PNG，"
                         f"请用 parse_regulation 工具读取其中的坐标与范围标注：")
            lines.extend(f"> - {p}" for p in rendered)
            print(f"[图件] {path.name}：{n_pages} 页中 {len(rendered)} 页含图，已渲染导出",
                  file=sys.stderr)

    return "\n".join(lines)


PARSERS = {
    ".pdf": ("PDF", parse_pdf),
    ".docx": ("Word", parse_docx),
    ".csv": ("Table", parse_csv),
    ".xlsx": ("Table", parse_excel),
    ".xlsm": ("Table", parse_excel),
}


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Word / Excel / CSV → 合并 Markdown（本地解析，不调用模型）")
    ap.add_argument("files", nargs="+", type=Path, help="输入文件（.docx/.csv/.xlsx）")
    ap.add_argument("-o", "--output", type=Path, help="输出路径，默认打印到标准输出")
    ap.add_argument("--image-dir", type=Path,
                    help="Word 内嵌图片的导出目录，默认与 -o 同目录；"
                         "导出的图片再交给 parse_regulation 解析")
    args = ap.parse_args()

    image_dir = args.image_dir or (args.output.parent if args.output else None)
    if image_dir is not None:
        image_dir.mkdir(parents=True, exist_ok=True)

    sections: List[str] = []
    for path in args.files:
        if not path.exists():
            print(f"[错误] 文件不存在：{path}", file=sys.stderr)
            return 1
        entry = PARSERS.get(path.suffix.lower())
        if entry is None:
            print(f"[错误] 不支持的格式 {path.suffix or '(无后缀)'}：{path.name}。"
                  f"支持 {'/'.join(PARSERS)}；PDF 与图片请用 parse_regulation 工具",
                  file=sys.stderr)
            return 1
        label, parser = entry
        try:
            sections.append(f"# [{label}] {path.name}\n\n{parser(path, image_dir)}")
        except Exception as exc:
            print(f"[错误] 解析失败 {path.name}：{exc}", file=sys.stderr)
            return 1

    result = "\n\n---\n\n".join(sections)
    if args.output:
        args.output.write_text(result, encoding="utf-8")
        print(f"[OK] 已解析 {len(sections)} 个文件，结果写入 {args.output}", file=sys.stderr)
    else:
        print(result)
    return 0


if __name__ == "__main__":
    sys.exit(main())
